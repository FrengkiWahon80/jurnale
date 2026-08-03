import streamlit as st
import sqlite3
from datetime import datetime
import io
import docx
import json

# --- 1. KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Jurnal SMPN Tujuh Maret Hadakewa", page_icon="🏫", layout="wide")

DEFAULT_SEKOLAH = "SMPN Tujuh Maret Hadakewa"
DB_NAME = "jurnal_smpn7_v2.db"  # Menggunakan versi DB baru agar tidak bentrok dengan data lama

# --- 2. DATABASE INITIALIZATION ---
def get_connection():
    conn = sqlite3.connect(DB_NAME, check_same_thread=False)
    conn.row_factory = sqlite3.Row  # Mengakses kolom berdasarkan NAMA, bukan angka indeks (Anti-Error)
    return conn

def init_db():
    conn = get_connection()
    c = conn.cursor()
    
    # Tabel Users
    c.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            password TEXT,
            nama_lengkap TEXT,
            nip TEXT,
            nama_sekolah TEXT DEFAULT 'SMPN Tujuh Maret Hadakewa',
            nama_kelas TEXT,
            role TEXT DEFAULT 'guru_wali'
        )
    ''')
    
    # Tabel Kepala Sekolah
    c.execute('''
        CREATE TABLE IF NOT EXISTS kepsek (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nama_kepsek TEXT,
            nip_kepsek TEXT,
            nama_sekolah TEXT DEFAULT 'SMPN Tujuh Maret Hadakewa'
        )
    ''')
    
    # Tabel Siswa
    c.execute('''
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            guru_id INTEGER,
            nama_siswa TEXT,
            nisn TEXT,
            nama_kelas TEXT,
            FOREIGN KEY(guru_id) REFERENCES users(id)
        )
    ''')
    
    # Tabel Master Kebiasaan SMP & 8 SKL
    c.execute('''
        CREATE TABLE IF NOT EXISTS habits (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nama_kebiasaan TEXT,
            kategori_skl TEXT
        )
    ''')
    
    # Tabel Jurnal Harian
    c.execute('''
        CREATE TABLE IF NOT EXISTS journal_entries (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_id INTEGER,
            tanggal TEXT,
            skor_kebiasaan TEXT,
            catatan_harian TEXT,
            FOREIGN KEY(student_id) REFERENCES students(id)
        )
    ''')
    
    # Seed Data Default Users jika Kosong
    c.execute("SELECT COUNT(*) FROM users")
    if c.fetchone()[0] == 0:
        c.execute("INSERT INTO users (username, password, nama_lengkap, nip, nama_sekolah, nama_kelas, role) VALUES (?, ?, ?, ?, ?, ?, ?)",
                  ('guru1', '12345', 'Budi Santoso, S.Pd.', '198501012010011001', DEFAULT_SEKOLAH, 'Kelas 7A', 'guru_wali'))
        c.execute("INSERT INTO users (username, password, nama_lengkap, nip, nama_sekolah, nama_kelas, role) VALUES (?, ?, ?, ?, ?, ?, ?)",
                  ('guru2', '12345', 'Siti Rahma, S.Pd.', '198802022011012002', DEFAULT_SEKOLAH, 'Kelas 7B', 'guru_wali'))
        c.execute("INSERT INTO users (username, password, nama_lengkap, nip, nama_sekolah, nama_kelas, role) VALUES (?, ?, ?, ?, ?, ?, ?)",
                  ('kepsek', '12345', 'Drs. Ahmad Dahlan, M.Pd.', '197003031995031001', DEFAULT_SEKOLAH, '-', 'kepsek'))

    # Seed Data Kepsek Default
    c.execute("SELECT COUNT(*) FROM kepsek")
    if c.fetchone()[0] == 0:
        c.execute("INSERT INTO kepsek (nama_kepsek, nip_kepsek, nama_sekolah) VALUES (?, ?, ?)",
                  ('Kepala Sekolah, M.Pd.', '197501012000121001', DEFAULT_SEKOLAH))

    # Seed Master Kebiasaan SMP
    c.execute("SELECT COUNT(*) FROM habits")
    if c.fetchone()[0] == 0:
        master_habits = [
            ("Beribadah Tepat Waktu & Berakhlak Mulia", "SKL Sikap Spiritual"),
            ("Bangun Pagi & Disiplin Diri", "SKL Sikap Sosial (Kemandirian)"),
            ("Berolahraga & Menjaga Kesehatan", "SKL Keterampilan (Jasmani & Kesehatan)"),
            ("Gemar Membaca / Literasi Buku", "SKL Pengetahuan & Literasi"),
            ("Gotong Royong & Kepedulian Lingkungan", "SKL Sikap Sosial (Empati)"),
            ("Belajar Mandiri & Pemanfaatan Teknologi", "SKL Pengetahuan & Proses Belajar"),
            ("Tidur Tepat Waktu & Istirahat Cukup", "SKL Sikap Sosial (Disiplin Diri)")
        ]
        c.executemany("INSERT INTO habits (nama_kebiasaan, kategori_skl) VALUES (?, ?)", master_habits)

    conn.commit()
    conn.close()

init_db()

# --- 3. AI CHARACTER SYNTHESIZER ---
def generate_ai_character_summary(nama_siswa, list_kebiasaan, list_catatan):
    total_jurnal = len(list_kebiasaan)
    if total_jurnal == 0:
        return f"Belum ada data jurnal harian yang cukup untuk dianalisis bagi siswa {nama_siswa}."

    skor_counts = {}
    for entry in list_kebiasaan:
        try:
            data = json.loads(entry)
            for k, v in data.items():
                if v:
                    skor_counts[k] = skor_counts.get(k, 0) + 1
        except Exception:
            pass

    summary = f"🤖 ANALISIS PERKEMBANGAN KARAKTER AI ({DEFAULT_SEKOLAH}):\n\n"
    summary += f"Berdasarkan rekapitulasi {total_jurnal} entri jurnal harian, ananda {nama_siswa} menunjukkan indikator perkembangan karakter sebagai berikut:\n\n"
    
    summary += "1. Sikap Spiritual & Sosial (SKL 1 & 2): "
    if skor_counts.get("Beribadah Tepat Waktu & Berakhlak Mulia", 0) / max(total_jurnal, 1) >= 0.7:
        summary += "Sangat baik dalam ketaatan beribadah dan menunjukkan tingkat kedisiplinan yang tinggi. "
    else:
        summary += "Sudah mulai menunjukkan ketaatan beribadah, namun masih perlu dorongan pembiasaan harian. "
        
    if skor_counts.get("Gotong Royong & Kepedulian Lingkungan", 0) / max(total_jurnal, 1) >= 0.6:
        summary += "Memiliki rasa empati dan kepedulian sosial yang sangat baik di lingkungan sekolah.\n"
    else:
        summary += "Sikap gotong royong dan kepedulian sosial berada pada tahap berkembang.\n"

    summary += "2. Literasi & Pembelajaran SMP (SKL Pengetahuan): "
    if skor_counts.get("Gemar Membaca / Literasi Buku", 0) / max(total_jurnal, 1) >= 0.6:
        summary += "Memiliki minat membaca dan kemandirian belajar yang kuat. "
    else:
        summary += "Perlu ditingkatkan motivasi literasi dan pemanfaatan waktu belajar mandiri. "

    summary += f"\n\n💡 Rekomendasi AI: Ananda {nama_siswa} disarankan untuk terus diapresiasi pada kebiasaan positifnya serta diberikan pendampingan pada aspek pembiasaan yang masih berkembang."
    
    return summary

# --- 4. SESSION STATE FOR AUTH ---
if 'user' not in st.session_state:
    st.session_state['user'] = None

# --- 5. TAMPILAN LOGIN & REGISTRASI ---
if st.session_state['user'] is None:
    st.title(f"🏫 Jurnal Guru - {DEFAULT_SEKOLAH}")
    
    auth_tab1, auth_tab2 = st.tabs(["🔐 Login Guru", "📝 Daftar Akun Guru Baru"])
    
    with auth_tab1:
        col1, _ = st.columns([1, 2])
        with col1:
            username_input = st.text_input("Username")
            password_input = st.text_input("Password", type="password")
            if st.button("Masuk Aplikasi", type="primary"):
                conn = get_connection()
                c = conn.cursor()
                c.execute("SELECT * FROM users WHERE username = ? AND password = ?", (username_input, password_input))
                user = c.fetchone()
                conn.close()
                
                if user:
                    st.session_state['user'] = {
                        'id': user['id'],
                        'username': user['username'],
                        'nama_lengkap': user['nama_lengkap'],
                        'nip': user['nip'],
                        'nama_sekolah': user['nama_sekolah'],
                        'nama_kelas': user['nama_kelas'],
                        'role': user['role']
                    }
                    st.success("Login Berhasil! Memuat data...")
                    st.rerun()
                else:
                    st.error("Username atau password salah!")

    with auth_tab2:
        st.markdown("##### Registrasi Akun Guru Wali Kelas Baru")
        reg_nama = st.text_input("Nama Lengkap Guru (dengan Gelar)")
        reg_nip = st.text_input("NIP Guru")
        reg_kelas = st.selectbox("Wali Kelas", ["Kelas 7A", "Kelas 7B", "Kelas 7C", "Kelas 8A", "Kelas 8B", "Kelas 9A", "Kelas 9B"])
        reg_username = st.text_input("Username Baru")
        reg_password = st.text_input("Password Baru", type="password")
        
        if st.button("Daftar Akun Baru", type="secondary"):
            if reg_nama and reg_username and reg_password:
                try:
                    conn = get_connection()
                    c = conn.cursor()
                    c.execute("INSERT INTO users (username, password, nama_lengkap, nip, nama_sekolah, nama_kelas, role) VALUES (?, ?, ?, ?, ?, ?, ?)",
                              (reg_username, reg_password, reg_nama, reg_nip, DEFAULT_SEKOLAH, reg_kelas, 'guru_wali'))
                    conn.commit()
                    conn.close()
                    st.success("Akun berhasil dibuat! Silakan Login di tab '🔐 Login Guru'.")
                except sqlite3.IntegrityError:
                    st.error("Username sudah digunakan. Silakan pilih username lain.")
            else:
                st.warning("Mohon isi semua data yang wajib!")

# --- 6. DASHBOARD UTAMA ---
else:
    user = st.session_state['user']
    st.sidebar.title(f"🏫 {DEFAULT_SEKOLAH}")
    st.sidebar.write(f"**Guru:** {user['nama_lengkap']}")
    st.sidebar.write(f"**NIP:** {user['nip'] if user['nip'] else '-'}")
    st.sidebar.write(f"**Wali Kelas:** {user['nama_kelas']}")
    
    if st.sidebar.button("Logout"):
        st.session_state['user'] = None
        st.rerun()

    st.title(f"🇮🇩 Control Panel Jurnal Kebiasaan - {user['nama_kelas']}")
    
    conn = get_connection()
    c = conn.cursor()

    # ISOLASI DATA GURU
    c.execute("SELECT * FROM students WHERE guru_id = ?", (user['id'],))
    students = c.fetchall()

    c.execute("SELECT * FROM habits")
    master_habits = c.fetchall()

    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "👨‍🎓 Kelola Data Siswa", 
        "📝 Jurnal Harian & 8 SKL", 
        "📊 Rekapitulasi Jurnal", 
        "🤖 AI Summary & Export Word",
        "⚙️ Profil Guru & Kepsek"
    ])

    # ==================== TAB 1: KELOLA DATA SISWA ====================
    with tab1:
        st.subheader(f"👨‍🎓 Data Siswa {user['nama_kelas']}")
        col_a, col_b = st.columns([1, 2])
        
        with col_a:
            st.markdown("##### Tambah Siswa Baru")
            nama_baru = st.text_input("Nama Lengkap Siswa")
            nisn_baru = st.text_input("NISN Siswa")

            if st.button("Tambah Siswa", type="primary"):
                if nama_baru and nisn_baru:
                    c.execute("INSERT INTO students (guru_id, nama_siswa, nisn, nama_kelas) VALUES (?, ?, ?, ?)", 
                              (user['id'], nama_baru, nisn_baru, user['nama_kelas']))
                    conn.commit()
                    st.success(f"Siswa {nama_baru} berhasil ditambahkan!")
                    st.rerun()
                else:
                    st.warning("Mohon isi Nama dan NISN Siswa!")

        with col_b:
            st.markdown(f"##### Daftar Siswa Kelas Anda ({len(students)} Siswa)")
            if students:
                for s in students:
                    col_s1, col_s2 = st.columns([3, 1])
                    col_s1.write(f"👤 **{s['nama_siswa']}** (NISN: {s['nisn']})")
                    if col_s2.button("Hapus", key=f"del_{s['id']}"):
                        c.execute("DELETE FROM students WHERE id = ? AND guru_id = ?", (s['id'], user['id']))
                        conn.commit()
                        st.rerun()
            else:
                st.info("Belum ada data siswa. Silakan tambah data di form sebelah kiri.")

    # ==================== TAB 2: INPUT JURNAL HARIAN & 8 SKL ====================
    with tab2:
        st.subheader("📝 Input Jurnal Harian & Catatan Wali Kelas")
        if students:
            student_dict = {f"{s['nama_siswa']} (NISN: {s['nisn']})": s['id'] for s in students}
            selected_student_label = st.selectbox("Pilih Siswa", list(student_dict.keys()), key="jurnal_select")
            selected_student_id = student_dict[selected_student_label]

            tanggal = st.date_input("Tanggal Input", datetime.now())
            
            st.markdown("##### Checklist Kebiasaan Indonesia Hebat (8 SKL SMP):")
            skor_kebiasaan_input = {}
            
            col_h1, col_h2 = st.columns(2)
            for i, h in enumerate(master_habits):
                col_target = col_h1 if i % 2 == 0 else col_h2
                skor_kebiasaan_input[h['nama_kebiasaan']] = col_target.checkbox(f"{h['nama_kebiasaan']} ({h['kategori_skl']})", value=True, key=f"habit_{h['id']}")

            catatan_harian = st.text_area("Catatan Perkembangan Harian Wali Kelas", placeholder="Tuliskan peristiwa atau perkembangan karakter siswa hari ini...")

            if st.button("Simpan Jurnal Harian", type="primary"):
                skor_json = json.dumps(skor_kebiasaan_input)
                c.execute("INSERT INTO journal_entries (student_id, tanggal, skor_kebiasaan, catatan_harian) VALUES (?, ?, ?, ?)",
                          (selected_student_id, tanggal.strftime('%Y-%m-%d'), skor_json, catatan_harian))
                conn.commit()
                st.success("Jurnal dan catatan harian berhasil disimpan!")
        else:
            st.warning("Silakan tambah data siswa Anda terlebih dahulu di Tab 'Kelola Data Siswa'.")

    # ==================== TAB 3: REKAPITULASI JURNAL ====================
    with tab3:
        st.subheader("📊 Rekapitulasi & Riwayat Jurnal")
        if students:
            student_dict_rekap = {f"{s['nama_siswa']} (NISN: {s['nisn']})": s['id'] for s in students}
            selected_student_rekap = st.selectbox("Pilih Siswa", list(student_dict_rekap.keys()), key="rekap_select")
            rekap_student_id = student_dict_rekap[selected_student_rekap]

            c.execute("SELECT tanggal, skor_kebiasaan, catatan_harian FROM journal_entries WHERE student_id = ? ORDER BY tanggal DESC", (rekap_student_id,))
            rekap_entries = c.fetchall()

            if rekap_entries:
                for entry in rekap_entries:
                    with st.expander(f"📅 Tanggal: {entry['tanggal']}"):
                        st.write(f"Catatan Harian Wali Kelas: {entry['catatan_harian'] if entry['catatan_harian'] else '-'}")
                        st.write("Capaian Kebiasaan:")
                        try:
                            data_skor = json.loads(entry['skor_kebiasaan'])
                            for k, v in data_skor.items():
                                status = "Ya" if v else "Tidak"
                                st.write(f"- {k}: {status}")
                        except Exception:
                            pass
            else:
                st.info("Belum ada riwayat jurnal untuk siswa ini.")

    # ==================== TAB 4: AI SUMMARY & EXPORT WORD ====================
    with tab4:
        st.subheader("🤖 AI Search & Generate Laporan Word (.docx)")
        if students:
            student_dict_ai = {f"{s['nama_siswa']} (NISN: {s['nisn']})": s['id'] for s in students}
            selected_student_ai = st.selectbox("Pilih Siswa untuk Export", list(student_dict_ai.keys()), key="ai_select")
            ai_student_id = student_dict_ai[selected_student_ai]

            c.execute("SELECT * FROM students WHERE id = ?", (ai_student_id,))
            s_data = c.fetchone()

            c.execute("SELECT skor_kebiasaan, catatan_harian, tanggal FROM journal_entries WHERE student_id = ?", (ai_student_id,))
            journal_data = c.fetchall()

            list_skor = [j['skor_kebiasaan'] for j in journal_data]
            list_catatan = [j['catatan_harian'] for j in journal_data]

            if st.button("Jalankan Analisis AI Karakter Siswa", type="secondary"):
                st.session_state['ai_result'] = generate_ai_character_summary(s_data['nama_siswa'], list_skor, list_catatan)

            ai_summary_text = st.text_area(
                "Kesimpulan AI & Catatan Akhir Wali Kelas (Dapat Diedit):", 
                value=st.session_state.get('ai_result', generate_ai_character_summary(s_data['nama_siswa'], list_skor, list_catatan)),
                height=200
            )

            c.execute("SELECT * FROM kepsek LIMIT 1")
            kepsek_data = c.fetchone()

            if st.button("Generate & Download Laporan MS Word", type="primary"):
                doc = docx.Document()
                
                doc.add_heading(f"{DEFAULT_SEKOLAH.upper()}", 0)
                doc.add_paragraph('LAPORAN JURNAL KEBIASAAN ANAK INDONESIA HEBAT (SMP)')
                doc.add_paragraph('TERINTEGRASI 8 STANDAR KOMPETENSI LULUSAN (SKL)')
                doc.add_paragraph('----------------------------------------------------------------------------------')
                
                doc.add_paragraph(f"Nama Siswa : {s_data['nama_siswa']}")
                doc.add_paragraph(f"NISN       : {s_data['nisn']}")
                doc.add_paragraph(f"Kelas      : {s_data['nama_kelas']}")
                doc.add_paragraph(f"Sekolah    : {DEFAULT_SEKOLAH}")
                doc.add_paragraph(f"Periode    : {datetime.now().strftime('%B %Y')}")

                doc.add_heading('1. KESIMPULAN PERKEMBANGAN KARAKTER (ANALISIS AI & WALI KELAS)', level=1)
                doc.add_paragraph(ai_summary_text)

                doc.add_heading('2. REKAP CATATAN HARIAN WALI KELAS', level=1)
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Tanggal'
                hdr_cells[1].text = 'Catatan Harian Perkembangan'

                if journal_data:
                    for j in journal_data:
                        if j['catatan_harian']:
                            row_cells = table.add_row().cells
                            row_cells[0].text = j['tanggal']
                            row_cells[1].text = j['catatan_harian']
                else:
                    row_cells = table.add_row().cells
                    row_cells[0].text = "-"
                    row_cells[1].text = "Belum ada catatan harian."

                doc.add_paragraph("\n\n")
                
                p_ttd = doc.add_paragraph()
                p_ttd.add_run(f"Mengetahui,\nKepala Sekolah {DEFAULT_SEKOLAH}\n\n\n\n")
                p_ttd.add_run(f"({kepsek_data['nama_kepsek'] if kepsek_data else 'Kepala Sekolah'})\n")
                p_ttd.add_run(f"NIP. {kepsek_data['nip_kepsek'] if kepsek_data else '-'}\t\t\t\t\t\t")
                
                p_ttd.add_run(f"Guru Wali Kelas {user['nama_kelas']}\n\n\n\n")
                p_ttd.add_run(f"({user['nama_lengkap']})\n")
                p_ttd.add_run(f"NIP. {user['nip'] if user['nip'] else '-'}")

                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)

                st.download_button(
                    label="📥 Klik Disini Untuk Unduh Laporan Word (.docx)",
                    data=bio,
                    file_name=f"Laporan_Jurnal_{s_data['nama_siswa'].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # ==================== TAB 5: PENGATURAN PROFIL ====================
    with tab5:
        st.subheader("⚙️ Pengaturan Identitas Guru & Kepala Sekolah")
        
        col_p1, col_p2 = st.columns(2)
        
        with col_p1:
            st.markdown("##### Identitas Guru Wali Kelas")
            prof_nama = st.text_input("Nama Lengkap Guru", value=user['nama_lengkap'])
            prof_nip = st.text_input("NIP Guru", value=user['nip'] if user['nip'] else '')
            prof_kelas = st.selectbox("Kelas Yang Diampu", ["Kelas 7A", "Kelas 7B", "Kelas 7C", "Kelas 8A", "Kelas 8B", "Kelas 9A", "Kelas 9B"], index=0)
            
            if st.button("Simpan Profil Guru", type="primary"):
                c.execute("UPDATE users SET nama_lengkap = ?, nip = ?, nama_kelas = ? WHERE id = ?",
                          (prof_nama, prof_nip, prof_kelas, user['id']))
                conn.commit()
                st.session_state['user']['nama_lengkap'] = prof_nama
                st.session_state['user']['nip'] = prof_nip
                st.session_state['user']['nama_kelas'] = prof_kelas
                st.success("Profil Guru berhasil diperbarui!")
                st.rerun()

        with col_p2:
            st.markdown("##### Identitas Kepala Sekolah")
            c.execute("SELECT * FROM kepsek LIMIT 1")
            k_data = c.fetchone()
            
            kepsek_nama = st.text_input("Nama Kepala Sekolah", value=k_data['nama_kepsek'] if k_data else '')
            kepsek_nip = st.text_input("NIP Kepala Sekolah", value=k_data['nip_kepsek'] if k_data else '')
            
            if st.button("Simpan Identitas Kepsek", type="secondary"):
                if k_data:
                    c.execute("UPDATE kepsek SET nama_kepsek = ?, nip_kepsek = ? WHERE id = ?", (kepsek_nama, kepsek_nip, k_data['id']))
                else:
                    c.execute("INSERT INTO kepsek (nama_kepsek, nip_kepsek, nama_sekolah) VALUES (?, ?, ?)", (kepsek_nama, kepsek_nip, DEFAULT_SEKOLAH))
                conn.commit()
                st.success("Identitas Kepala Sekolah berhasil disimpan!")
                st.rerun()

    conn.close()
